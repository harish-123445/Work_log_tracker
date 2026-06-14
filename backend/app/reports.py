import io
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListFlowable,
    ListItem,
    Table,
    TableStyle,
    HRFlowable,
)

from app import schemas


def _fmt_date(d):
    return d.strftime("%b %Y") if d else "Present" if d is None else ""


def _date_range(project: "schemas.Project") -> str:
    start = project.start_date.strftime("%b %Y") if project.start_date else "—"
    if project.status == "Ongoing" and not project.end_date:
        end = "Present"
    else:
        end = project.end_date.strftime("%b %Y") if project.end_date else "—"
    return f"{start} – {end}"


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _add_project_to_doc(doc: Document, project: "schemas.Project"):
    heading = doc.add_heading(level=2)
    run = heading.add_run(project.project_title)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x3D)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"{project.company_name}"
        + (f"  •  {project.role_title}" if project.role_title else "")
        + f"  •  {_date_range(project)}  •  {project.status}"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x66, 0x70, 0x6B)

    if project.project_manager_name:
        pm = doc.add_paragraph()
        pm_run = pm.add_run(
            "Project Manager: "
            + project.project_manager_name
            + (f" ({project.project_manager_contact})" if project.project_manager_contact else "")
        )
        pm_run.font.size = Pt(10)

    if project.description:
        doc.add_paragraph(project.description)

    if project.features:
        doc.add_paragraph("Key Features / Responsibilities:", style="Intense Quote")
        for feature in project.features:
            doc.add_paragraph(feature, style="List Bullet")

    if project.technologies:
        tech_para = doc.add_paragraph()
        tech_run = tech_para.add_run("Technologies: " + ", ".join(project.technologies))
        tech_run.font.size = Pt(10)

    if project.achievements:
        doc.add_paragraph("Achievements / Impact:", style="Intense Quote")
        doc.add_paragraph(project.achievements)

    if project.notes:
        doc.add_paragraph("Notes:", style="Intense Quote")
        doc.add_paragraph(project.notes)

    doc.add_paragraph()  # spacer


def build_project_docx(project: "schemas.Project") -> io.BytesIO:
    doc = Document()
    _add_project_to_doc(doc, project)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_full_docx(user: "schemas.UserOut", projects: list["schemas.Project"]) -> io.BytesIO:
    doc = Document()

    title = doc.add_heading(level=0)
    title_run = title.add_run("Work Log Report")
    title_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x3D)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(
        f"{user.full_name or user.username}  •  Generated on {date.today().strftime('%d %b %Y')}"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)

    doc.add_paragraph()

    # Group projects by company
    companies: dict[str, list] = {}
    for p in projects:
        companies.setdefault(p.company_name, []).append(p)

    for company, comp_projects in companies.items():
        doc.add_heading(company, level=1)
        for project in comp_projects:
            _add_project_to_doc(doc, project)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

INK = colors.HexColor("#1F3A3D")
MUTED = colors.HexColor("#6B7570")
ACCENT = colors.HexColor("#D98E48")


def _pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("ReportTitle", parent=styles["Title"], textColor=INK, fontSize=22))
    styles.add(ParagraphStyle("CompanyHeading", parent=styles["Heading1"], textColor=INK, spaceBefore=18))
    styles.add(ParagraphStyle("ProjectHeading", parent=styles["Heading2"], textColor=INK, spaceBefore=10))
    styles.add(ParagraphStyle("Meta", parent=styles["Normal"], textColor=MUTED, fontSize=9, spaceAfter=4))
    styles.add(ParagraphStyle("BodyTextSmall", parent=styles["BodyText"], fontSize=10, spaceAfter=6))
    styles.add(ParagraphStyle("SectionLabel", parent=styles["Normal"], textColor=ACCENT, fontSize=10, fontName="Helvetica-Bold", spaceBefore=6, spaceAfter=2))
    return styles


def _project_flowables(project: "schemas.Project", styles) -> list:
    flow = []
    flow.append(Paragraph(project.project_title, styles["ProjectHeading"]))

    meta_bits = [project.company_name]
    if project.role_title:
        meta_bits.append(project.role_title)
    meta_bits.append(_date_range(project))
    meta_bits.append(project.status)
    flow.append(Paragraph("  •  ".join(meta_bits), styles["Meta"]))

    if project.project_manager_name:
        pm_text = f"Project Manager: {project.project_manager_name}"
        if project.project_manager_contact:
            pm_text += f" ({project.project_manager_contact})"
        flow.append(Paragraph(pm_text, styles["Meta"]))

    if project.description:
        flow.append(Paragraph(project.description, styles["BodyTextSmall"]))

    if project.features:
        flow.append(Paragraph("KEY FEATURES / RESPONSIBILITIES", styles["SectionLabel"]))
        flow.append(
            ListFlowable(
                [ListItem(Paragraph(f, styles["BodyTextSmall"])) for f in project.features],
                bulletType="bullet",
                leftIndent=14,
            )
        )

    if project.technologies:
        flow.append(Paragraph("TECHNOLOGIES: " + ", ".join(project.technologies), styles["Meta"]))

    if project.achievements:
        flow.append(Paragraph("ACHIEVEMENTS / IMPACT", styles["SectionLabel"]))
        flow.append(Paragraph(project.achievements, styles["BodyTextSmall"]))

    if project.notes:
        flow.append(Paragraph("NOTES", styles["SectionLabel"]))
        flow.append(Paragraph(project.notes, styles["BodyTextSmall"]))

    flow.append(Spacer(1, 8))
    flow.append(HRFlowable(width="100%", color=colors.HexColor("#E0E0E0"), thickness=0.5))
    flow.append(Spacer(1, 8))
    return flow


def build_project_pdf(project: "schemas.Project") -> io.BytesIO:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = _pdf_styles()

    elements = [Paragraph(project.project_title, styles["ReportTitle"]), Spacer(1, 10)]
    elements += _project_flowables(project, styles)

    doc.build(elements)
    buf.seek(0)
    return buf


def build_full_pdf(user: "schemas.UserOut", projects: list["schemas.Project"]) -> io.BytesIO:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = _pdf_styles()

    elements = [
        Paragraph("Work Log Report", styles["ReportTitle"]),
        Paragraph(
            f"{user.full_name or user.username}  •  Generated on {date.today().strftime('%d %b %Y')}",
            styles["Meta"],
        ),
        Spacer(1, 12),
    ]

    companies: dict[str, list] = {}
    for p in projects:
        companies.setdefault(p.company_name, []).append(p)

    for company, comp_projects in companies.items():
        elements.append(Paragraph(company, styles["CompanyHeading"]))
        for project in comp_projects:
            elements += _project_flowables(project, styles)

    doc.build(elements)
    buf.seek(0)
    return buf
